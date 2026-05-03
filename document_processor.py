import fitz  # pymupdf
from docx import Document
from io import BytesIO
from typing import List, Dict


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            if break_point > chunk_size // 2:
                chunk = chunk[:break_point + 1]
        chunk = chunk.strip()
        if chunk:
            chunks.append(chunk)
        start += max(len(chunk) - overlap, 1)
    return chunks


def process_pdf(file_bytes: bytes) -> List[Dict]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            pages.append({"page": page_num, "text": text})
    doc.close()
    return pages


def process_docx(file_bytes: bytes) -> List[Dict]:
    doc = Document(BytesIO(file_bytes))
    paragraphs_by_page: List[str] = []
    current_block: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Treat each heading as a page boundary for citation purposes
        if para.style.name.startswith("Heading"):
            if current_block:
                paragraphs_by_page.append("\n".join(current_block))
                current_block = []
        current_block.append(text)

    if current_block:
        paragraphs_by_page.append("\n".join(current_block))

    # If no headings found, treat the whole doc as one "page"
    if not paragraphs_by_page:
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        paragraphs_by_page = [full_text] if full_text else []

    return [{"page": i + 1, "text": text} for i, text in enumerate(paragraphs_by_page)]


def extract_chunks(file_bytes: bytes, filename: str) -> List[Dict]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        pages = process_pdf(file_bytes)
    elif lower.endswith(".docx"):
        pages = process_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    chunks: List[Dict] = []
    chunk_idx = 0
    for page_data in pages:
        for chunk_text_val in chunk_text(page_data["text"]):
            chunks.append({
                "text": chunk_text_val,
                "page": page_data["page"],
                "chunk_index": chunk_idx,
            })
            chunk_idx += 1

    return chunks
